from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any, Iterable, Optional

log = logging.getLogger("meeting")

def write_text_docx(path: Path, title: str, paragraphs: Iterable[str]) -> None:
    """Write a small DOCX document from a title and plain-text paragraphs."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)
    for paragraph in paragraphs:
        text = str(paragraph).strip()
        if text:
            doc.add_paragraph(text)
    doc.save(path)
    log.info("Saved DOCX: %s", path)

def _docx_text(value: Any) -> str:
    if isinstance(value, dict):
        return value.get("text") or value.get("name") or json.dumps(
            value, ensure_ascii=False
        )
    return str(value)

def write_summary_docx(
    title: str,
    speaker: str,
    duration: str,
    topics: Iterable[Any],
    path: Path,
    key_concepts: Optional[Iterable[Any]] = None,
    language: str = "en",
) -> None:
    """Write a meeting or lecture summary with topics and timestamped concepts."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)
    russian = language.lower().startswith("ru")
    speaker_label = "Участники" if russian else "Speaker"
    duration_label = "Продолжительность" if russian else "Duration"
    topics_label = "Ключевые темы" if russian else "Key topics"
    concepts_label = "Ключевые выводы" if russian else "Key concepts"
    metadata = "; ".join(
        value for value in (f"{speaker_label}: {speaker}" if speaker else "", f"{duration_label}: {duration}" if duration else "") if value
    )
    if metadata:
        doc.add_paragraph(metadata)
    doc.add_heading(topics_label, level=1)
    for topic in topics:
        if isinstance(topic, dict) and topic.get("heading"):
            doc.add_heading(str(topic["heading"]), level=2)
            doc.add_paragraph(str(topic.get("text", "")))
        else:
            doc.add_paragraph(_docx_text(topic), style="List Bullet")
    concepts = list(key_concepts or [])
    if concepts:
        doc.add_heading(concepts_label, level=1)
        for concept in concepts:
            if isinstance(concept, dict):
                timestamp = concept.get("timestamp") or concept.get("time")
                text = concept.get("text") or concept.get("concept") or _docx_text(concept)
                doc.add_paragraph(
                    f"{timestamp}: {text}" if timestamp else text,
                    style="List Bullet",
                )
            else:
                doc.add_paragraph(str(concept), style="List Bullet")
    doc.save(path)
    log.info("Saved DOCX: %s", path)

def write_analytical_docx(sections: dict[str, str], path: Path, language: str = "en") -> None:
    """Write an analytical note with the supplied section text."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Аналитическая записка" if language.lower().startswith("ru") else "Analytical note", level=0)
    for title, content in sections.items():
        doc.add_heading(str(title), level=1)
        for paragraph in str(content).split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    doc.save(path)
    log.info("Saved DOCX: %s", path)

def write_protocol_docx(protocol: dict, path: Path) -> None:
    if protocol.get("quality", {}).get("valid") is False:
        return
    from docx import Document
    from docx.shared import Pt, Cm

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    doc.add_heading("Протокол совещания", level=0)
    doc.add_paragraph("")

    if protocol.get("participants"):
        doc.add_heading("Участники", level=1)
        table = doc.add_table(rows=1 + len(protocol["participants"]), cols=3)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "№"
        table.rows[0].cells[1].text = "Участник"
        table.rows[0].cells[2].text = "Основание (стенограмма)"
        for idx, p in enumerate(protocol["participants"], 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = p.get("name", str(p))
            table.rows[idx].cells[2].text = p.get("source_quote", "")

    for section_name in [
        "agenda",
        "decisions",
        "assignments",
        "open_questions",
        "risks",
        "next_steps",
    ]:
        items = protocol.get(section_name, [])
        if not items:
            continue
        section_titles = {
            "agenda": "Повестка",
            "decisions": "Решения",
            "assignments": "Поручения",
            "open_questions": "Открытые вопросы",
            "risks": "Риски",
            "next_steps": "Следующие шаги",
        }
        doc.add_heading(section_titles[section_name], level=1)
        for idx, item in enumerate(items, 1):
            text = (
                item.get("text")
                or item.get("task")
                or item.get("action")
                or item.get("question")
                or str(item)
            )
            doc.add_paragraph(f"{idx}. {text}")
            if item.get("assignee") or item.get("who") or item.get("deadline") or item.get("when"):
                assignee = item.get("assignee") or item.get("who") or "не указан"
                deadline = item.get("deadline") or item.get("when") or "не указан"
                doc.add_paragraph(f"Ответственный: {assignee}; срок: {deadline}.")
            if item.get("source_quote"):
                doc.add_paragraph(f"Основание (стенограмма): «{item['source_quote']}»")
    quality = protocol.get("quality", {})
    if quality.get("warnings"):
        doc.add_heading("Предупреждения о качестве", level=1)
        for warning in quality["warnings"]:
            doc.add_paragraph(str(warning), style="List Bullet")
    doc.save(path)
    log.info("Saved DOCX: %s", path)
