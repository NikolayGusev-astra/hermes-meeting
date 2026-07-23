#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
from importlib.util import find_spec
import json
import logging
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any, Iterable, List, Optional, Tuple
from urllib.parse import urlparse

from tenacity import retry, stop_after_attempt, wait_exponential

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s"
)
log = logging.getLogger("meeting")

MAX_FILE_MB = int(os.getenv("MEETING_MAX_FILE_MB", "2048"))
MAX_DURATION_SEC = int(os.getenv("MEETING_MAX_DURATION_SEC", "7200"))


def _transcribe_default_device() -> str:
    """Auto-detect best available device: cuda > mps > rocm > cpu."""
    import platform as _platform

    # 1. NVIDIA CUDA
    try:
        from ctranslate2 import get_cuda_device_count as _cuda_count

        if _cuda_count() > 0:
            # Probe for CUDA runtime: check pip-installed nvidia-cublas-cu12 or system PATH
            _cublas_ok = False
            import ctypes as _ct

            for _dll in ("cublas64_12.dll", "cublas64_11.dll", "libcublas.so.12"):
                try:
                    _ct.cdll.LoadLibrary(_dll)
                    _cublas_ok = True
                    break
                except OSError:
                    pass
            if not _cublas_ok:
                try:
                    import nvidia.cublas
                    _cublas_dir = Path(nvidia.cublas.__path__[0]) / "bin" / "cublas64_12.dll"
                    _ct.cdll.LoadLibrary(str(_cublas_dir))
                    _cublas_ok = True
                except Exception:
                    pass
            if not _cublas_ok:
                log.warning("CUDA GPU found but runtime missing. Install: pip install meeting-intelligence[gpu]")
                return "cpu"
            log.info("Auto-detected device: cuda")
            return "cuda"
    except Exception:
        pass

    if _platform.system() == "Darwin" and _platform.machine() == "arm64":
        log.info("Auto-detected device: cpu (Apple Silicon)")
        return "cpu"

    return "cpu"


TRANSCRIBE_MODEL = os.getenv("MEETING_TRANSCRIBE_MODEL") or (
    "large-v3-turbo" if _transcribe_default_device() == "cuda" else "small"
)
TRANSCRIBE_DEVICE = os.getenv("MEETING_TRANSCRIBE_DEVICE") or _transcribe_default_device()
TRANSCRIBE_COMPUTE = os.getenv("MEETING_TRANSCRIBE_COMPUTE") or (
    "float16" if TRANSCRIBE_DEVICE == "cuda" else "int8"
)
TRANSCRIBE_LANG = os.getenv("MEETING_TRANSCRIBE_LANG", "en")
LLM_BASE_URL = os.getenv("MEETING_LLM_BASE_URL", "http://localhost:1234/v1")
LLM_API_KEY = os.getenv("MEETING_LLM_API_KEY", "lm-studio")
LLM_MODEL = os.getenv("MEETING_LLM_MODEL", "qwen2.5-7b-instruct")
TRANSLATE_BATCH_SIZE = int(os.getenv("MEETING_TRANSLATE_BATCH_SIZE", "8"))


PROTOCOL_CHUNK_THRESHOLD_TOKENS = 6000
PROTOCOL_CHARS_PER_TOKEN = 4
PROTOCOL_SECTIONS = (
    "participants",
    "agenda",
    "decisions",
    "assignments",
    "open_questions",
    "unclear",
)

RUSSIAN_TRANSCRIPTION_PATTERNS = re.compile(
    r"\b(?:"
    r"akhmetov|yakovlev|ivanovich|petrovich|sergeevich|alexandrovich|"
    r"vladimirovich|ovna|evna|ichna|"
    r"minjust|minfin|minzdrav|minpromtorg|minobrnauki|mincifry|"
    r"rosstandart|rostandart|rospotrebnadzor|rosreestr|roskomnadzor|"
    r"gosduma|sovfed|pravitelstvo|"
    r"russian federation"
    r")\b",
    re.IGNORECASE,
)


def _is_probably_russian_mistranscribed_as_english(
    transcript: str, detected_language: str
) -> bool:
    """Identify English transcripts with a concentrated set of Russian cues."""
    if detected_language.lower() != "en":
        return False

    word_count = len(re.findall(r"\b\w+\b", transcript))
    russian_cue_count = len(RUSSIAN_TRANSCRIPTION_PATTERNS.findall(transcript))
    return russian_cue_count >= 2 and russian_cue_count / max(word_count, 1) >= 0.01


class MeetingError(Exception):
    pass


def fail(message: str, code: int = 2) -> None:
    log.error(message)
    raise SystemExit(code)


def _handle_exception(exc: Exception) -> None:
    msg = str(exc)
    log.error("Meeting pipeline error: %s", msg)
    raise MeetingError(msg) from exc


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def _is_url(value: str | Path) -> bool:
    parsed = urlparse(str(value))
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def _resolve_source(url_or_path: str | Path) -> Path:
    if not _is_url(url_or_path):
        return Path(url_or_path)
    if find_spec("yt_dlp") is None:
        fail("URL support requires yt-dlp; install meeting-intelligence[url]")

    download_dir = Path(tempfile.mkdtemp(prefix="meeting-intelligence-"))
    output_template = download_dir / "audio.%(ext)s"
    command = [
        sys.executable,
        "-m",
        "yt_dlp",
        "-x",
        "--audio-format",
        "wav",
        "--no-playlist",
        "--proxy", os.getenv("MEETING_YT_PROXY", ""),
        "--output",
        str(output_template),
        str(url_or_path),
    ]
    log.info("Downloading audio from URL")
    try:
        result = subprocess.run(
            command, capture_output=True, text=True, timeout=600, check=False
        )
    except subprocess.TimeoutExpired as exc:
        raise MeetingError("yt-dlp timed out while downloading audio") from exc
    if result.returncode != 0:
        raise MeetingError(f"yt-dlp failed: {result.stderr[-400:]}")

    audio = download_dir / "audio.wav"
    if not audio.is_file() or audio.stat().st_size == 0:
        raise MeetingError("yt-dlp did not produce a WAV audio file")
    return audio


def is_loopback_url(url: str) -> bool:
    host = (urlparse(url).hostname or "").lower()
    return host in {"127.0.0.1", "localhost", "::1", ""}


def enforce_cloud_policy(allow_cloud: bool) -> None:
    if not allow_cloud and not is_loopback_url(LLM_BASE_URL):
        fail(
            f"Cloud LLM is disabled; external URL {LLM_BASE_URL!r}. "
            f"Pass --allow-cloud explicitly."
        )


def _probe_duration(path: Path) -> float:
    try:
        proc = subprocess.run(
            ["ffprobe", "-v", "quiet", "-show_format", "-of", "json", str(path)],
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
        )
        if proc.returncode != 0:
            raise MeetingError(f"ffprobe failed: {proc.stderr[-200:]}")
        payload = json.loads(proc.stdout or "{}")
        duration = payload.get("format", {}).get("duration")
        if duration is None:
            raise MeetingError("ffprobe did not return duration")
        return float(duration)
    except Exception as exc:
        raise MeetingError(f"Duration probe failed: {exc}") from exc


def check_resource_limits(
    path: Path,
    *,
    max_file_mb: Optional[int] = None,
    max_duration_sec: Optional[int] = None,
) -> None:
    size_mb = path.stat().st_size / (1024 * 1024)
    if max_file_mb is None:
        max_file_mb = int(os.getenv("MEETING_MAX_FILE_MB", "2048"))
    if max_duration_sec is None:
        max_duration_sec = int(os.getenv("MEETING_MAX_DURATION_SEC", "7200"))
    if size_mb > max_file_mb:
        fail(f"File too large: {size_mb:.1f} MB > {max_file_mb} MB")
    duration = _probe_duration(path)
    if duration > max_duration_sec:
        fail(f"Duration too long: {duration:.0f}s > {max_duration_sec}s")


def extract_audio(src: Path, dst: Path) -> Path:
    dst.parent.mkdir(parents=True, exist_ok=True)
    cmd = [
        "ffmpeg",
        "-y",
        "-i",
        str(src),
        "-vn",
        "-acodec",
        "pcm_s16le",
        "-ar",
        "16000",
        "-ac",
        "1",
        str(dst),
    ]
    log.info("Extracting audio -> %s", dst)
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=600, check=False)
    if res.returncode != 0:
        fail(f"ffmpeg failed: {res.stderr[-400:]}")
    if not dst.exists() or dst.stat().st_size == 0:
        fail("ffmpeg produced empty audio file")
    return dst


def _silence_speakers(segments, silence_gap: float = 1.5):
    current = 0
    previous_end = None
    out = []
    for seg in segments:
        start = float(seg["start"])
        if previous_end is not None and (start - previous_end) > silence_gap:
            current += 1
        item = dict(seg)
        item["speaker_id"] = f"SPEAKER_{current:02d}"
        out.append(item)
        previous_end = float(item["end"])
    return out


def transcribe_audio(
    audio: Path, model: str, language: Optional[str], device: str, compute_type: str
) -> Tuple[str, dict]:
    from faster_whisper import WhisperModel

    log.info(
        "Loading whisper model=%s device=%s compute_type=%s",
        model,
        device,
        compute_type,
    )
    m = WhisperModel(model, device=device, compute_type=compute_type)
    segments_iter, info = m.transcribe(
        str(audio),
        language=language,
        beam_size=5,
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500),
    )
    log.info("Detected language=%s duration=%.1fs", info.language, info.duration)
    segments = []
    for idx, seg in enumerate(segments_iter, 1):
        text = seg.text.strip()
        if not text:
            continue
        item = {
            "id": f"seg_{idx:04d}",
            "start": float(seg.start),
            "end": float(seg.end),
            "text": text,
            "speaker_id": "SPEAKER_00",
        }
        segments.append(item)
        start_min = int(seg.start // 60)
        start_sec = int(seg.start % 60)
        end_min = int(seg.end // 60)
        end_sec = int(seg.end % 60)
        item["timestamp"] = (
            f"[{start_min:02d}:{start_sec:02d}->{end_min:02d}:{end_sec:02d}]"
        )
    garbage_runs = []
    run_start = None
    for index, item in enumerate(segments):
        if len(item["text"]) <= 2:
            if run_start is None:
                run_start = index
        elif run_start is not None:
            if index - run_start >= 5:
                garbage_runs.append((run_start, index))
            run_start = None
    if run_start is not None and len(segments) - run_start >= 5:
        garbage_runs.append((run_start, len(segments)))

    if garbage_runs:
        garbage_indexes = {
            index for start, end in garbage_runs for index in range(start, end)
        }
        log.warning("Removed %d short Whisper artifact segments", len(garbage_indexes))
        segments = [
            item for index, item in enumerate(segments) if index not in garbage_indexes
        ]

    enriched = _silence_speakers(segments)
    final_lines = []
    for item in enriched:
        final_lines.append(f"{item['timestamp']} {item['speaker_id']} | {item['text']}")
    transcript = "\n".join(final_lines)
    if _is_probably_russian_mistranscribed_as_english(transcript, info.language):
        log.warning(
            "Transcript may be Russian misdetected as English. "
            "Re-run with --language ru for better quality."
        )
    meta = {
        "schema_version": "0.1.0",
        "stt_model": model,
        "language": info.language,
        "language_probability": float(
            getattr(info, "language_probability", 0.0) or 0.0
        ),
        "no_speech_prob": float(getattr(info, "no_speech_prob", 0.0) or 0.0),
        "duration": float(info.duration),
        "segment_count": len(enriched),
    }
    return transcript, meta


def _clean_whisper_artifacts(transcript: str) -> str:
    """Remove Whisper hallucination lines made of repeated one-character tokens."""
    clean_lines = []
    for line in transcript.splitlines():
        single_char_tokens = [token for token in line.split() if len(token) == 1]
        if (
            len(single_char_tokens) >= 4
            and len(set(single_char_tokens)) == 1
        ):
            log.warning("Removed repeated-token Whisper artifact line")
            continue
        clean_lines.append(line)
    return "\n".join(clean_lines)


_SEGMENT_ID_PREFIX = re.compile(
    r"^\s*(?:\[?(?:seg(?:ment)?[_ -]?\d+)\]?:?\s*(?:SPEAKER_\d+\s*)?)",
    re.IGNORECASE,
)


def _agent_mode_enabled() -> bool:
    return os.getenv("MEETING_AGENT_MODE", "false").lower() in {"1", "true", "yes", "on"}


def prepare_agent_transcript(transcript: str, source: Path) -> dict[str, Any]:
    """Build the LLM-free, cleaned transcript payload consumed by meeting agents."""
    source_lines = transcript.splitlines()
    without_artifacts = _clean_whisper_artifacts(transcript).splitlines()
    cleaned_lines = []
    segment_ids_stripped = 0
    for line in without_artifacts:
        cleaned = _SEGMENT_ID_PREFIX.sub("", line).strip()
        if cleaned != line.strip():
            segment_ids_stripped += 1
        if cleaned:
            cleaned_lines.append(cleaned)
    cleaned_transcript = "\n".join(cleaned_lines)
    return {
        "schema_version": "0.7.0",
        "transcript": cleaned_transcript,
        "metadata": {
            "source": str(source),
            "source_hash": sha256(source),
            "input_line_count": len(source_lines),
            "output_line_count": len(cleaned_lines),
            "garbage_lines_removed": len(source_lines) - len(without_artifacts),
            "segment_ids_stripped": segment_ids_stripped,
            "llm_called": False,
        },
    }


def cmd_agent_transcript(args: argparse.Namespace) -> int:
    """Emit a cleaned transcript JSON payload for agent consumption without an LLM call."""
    src = Path(args.transcript)
    if not src.exists():
        fail(f"Transcript not found: {src}")
    payload = prepare_agent_transcript(src.read_text(encoding="utf-8"), src)
    if getattr(args, "docx", False):
        output = (
            Path(args.output)
            if getattr(args, "output", None)
            else src.with_suffix(".agent-transcript.docx")
        )
        write_text_docx(output, src.stem, payload["transcript"].splitlines())
    print(json.dumps(payload))
    return 0


def _read_docx_input(path: Path) -> tuple[Optional[dict[str, Any]], str]:
    text = path.read_text(encoding="utf-8")
    try:
        payload = json.loads(text)
    except json.JSONDecodeError:
        return None, text
    if not isinstance(payload, dict):
        fail("DOCX JSON input must be an object")
    return payload, text


def cmd_generate_docx(args: argparse.Namespace) -> int:
    """Generate summary, analytical, or protocol DOCX from JSON or text input."""
    src = Path(args.input)
    if not src.exists():
        fail(f"Input not found: {src}")
    payload, text = _read_docx_input(src)
    output = Path(args.output)

    if args.type == "summary":
        if payload:
            write_summary_docx(
                str(payload.get("title", src.stem)),
                str(payload.get("speaker", "")),
                str(payload.get("duration", "")),
                payload.get("topics", []),
                output,
                payload.get("key_concepts", payload.get("concepts", [])),
            )
        else:
            write_text_docx(output, src.stem, text.splitlines())
    elif args.type == "analytical":
        if payload:
            sections = payload.get("sections", payload)
            if not isinstance(sections, dict):
                fail("Analytical JSON input must contain an object of sections")
        else:
            sections = {"Context": text}
        write_analytical_docx(sections, output)
    else:
        if payload:
            write_protocol_docx(payload, output)
        else:
            write_text_docx(output, "Meeting protocol", text.splitlines())
    return 0


@retry(stop=stop_after_attempt(3), wait=wait_exponential(min=1, max=10))
def _translate_one(client: Any, text: str, target_lang: str) -> str:
    prompt = f"Translate to {target_lang}. Keep names/codes/technical terms unchanged. Output ONLY translation, no extra text.\n\n{text}"
    return (
        client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
        )
        .choices[0]
        .message.content.strip()
    )


def _chunked(seq: List[str], size: int) -> Iterable[List[str]]:
    for i in range(0, len(seq), size):
        yield seq[i : i + size]


def translate_lines(lines: List[str], target_lang: str, allow_cloud: bool) -> List[str]:
    enforce_cloud_policy(allow_cloud)
    if not lines:
        return []
    from openai import OpenAI

    client = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)
    out: List[str] = []
    failed = 0
    head_pat = re.compile(
        r"^\[(\d{2}:\d{2})->(\d{2}:\d{2})\]\s+(SPEAKER_\d+)\s+\|\s+(.*)$"
    )
    chunk_buf = []
    prefix_buf = []
    for line in lines:
        m = head_pat.match(line)
        if m:
            prefix_buf.append(f"[{m.group(1)}->{m.group(2)}] {m.group(3)} | ")
            chunk_buf.append(m.group(4))
        else:
            prefix_buf.append("")
            chunk_buf.append(line)
        if len(chunk_buf) >= TRANSLATE_BATCH_SIZE:
            try:
                translated = _translate_one(client, "\n".join(chunk_buf), target_lang)
                parts = translated.splitlines()
                if len(parts) != len(chunk_buf):
                    raise ValueError("Translator returned wrong line count")
                out.extend(
                    f"{p}{part}"
                    for p, part in zip(prefix_buf[-len(chunk_buf) :], parts)
                )
            except Exception as exc:
                failed += len(chunk_buf)
                log.warning("Batch translate failed: %s", exc)
            finally:
                chunk_buf.clear()
                prefix_buf.clear()
    if chunk_buf:
        try:
            translated = _translate_one(client, "\n".join(chunk_buf), target_lang)
            parts = translated.splitlines()
            if len(parts) != len(chunk_buf):
                raise ValueError("Translator returned wrong line count")
            out.extend(f"{p}{part}" for p, part in zip(prefix_buf, parts))
        except Exception as exc:
            failed += len(chunk_buf)
            log.warning("Batch translate failed: %s", exc)
    if failed:
        fail(f"LLM translation failed for {failed} line(s)")
    return out


def _protocol_chunk_size_tokens() -> int:
    value = os.getenv(
        "MEETING_PROTOCOL_CHUNK_SIZE", str(PROTOCOL_CHUNK_THRESHOLD_TOKENS)
    )
    try:
        return max(1, int(value))
    except ValueError:
        log.warning(
            "Invalid MEETING_PROTOCOL_CHUNK_SIZE=%r; using %s tokens",
            value,
            PROTOCOL_CHUNK_THRESHOLD_TOKENS,
        )
        return PROTOCOL_CHUNK_THRESHOLD_TOKENS


def _needs_protocol_chunking(transcript: str) -> bool:
    return len(transcript) > (
        PROTOCOL_CHUNK_THRESHOLD_TOKENS * PROTOCOL_CHARS_PER_TOKEN
    )


def _split_protocol_transcript(transcript: str, chunk_size_chars: int) -> List[str]:
    overlap_chars = max(1, chunk_size_chars // 10)
    chunks = []
    start = 0
    while start < len(transcript):
        end = min(len(transcript), start + chunk_size_chars)
        if end < len(transcript):
            line_end = transcript.rfind("\n", start + 1, end + 1)
            if line_end > start:
                end = line_end + 1
        chunks.append(transcript[start:end])
        if end == len(transcript):
            break
        start = max(start + 1, end - overlap_chars)
    return chunks


def _merge_protocol_chunks(protocols: Iterable[dict]) -> dict:
    merged = {section: [] for section in PROTOCOL_SECTIONS}
    seen_quotes = {section: set() for section in PROTOCOL_SECTIONS}
    for protocol in protocols:
        for section in PROTOCOL_SECTIONS:
            items = protocol.get(section, [])
            if not isinstance(items, list):
                continue
            for item in items:
                quote = (
                    _normalize(item.get("source_quote", ""))
                    if isinstance(item, dict)
                    else ""
                )
                if quote and quote in seen_quotes[section]:
                    continue
                if quote:
                    seen_quotes[section].add(quote)
                merged[section].append(item)
    return merged


def _repair_json(text: str):
    """Try to fix common LLM JSON errors: single quotes, trailing commas."""
    import re

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    cleaned = re.sub(r"```(?:json)?\s*", "", text).strip()
    cleaned = re.sub(r"'([^']*)'(\s*:)", r'"\1"\2', cleaned)
    cleaned = re.sub(r":\s*'([^']*)'", r': "\1"', cleaned)
    cleaned = re.sub(r",(\s*[}\]])", r"\1", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        return None


def _build_protocol_chunk(transcript: str, model: str, allow_cloud: bool) -> dict:
    enforce_cloud_policy(allow_cloud)
    from openai import OpenAI

    client = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)
    # Strip segment IDs and duplicate speaker labels for cleaner LLM input
    import re as _re

    clean_transcript = _re.sub(
        r"^seg_\d+\s+SPEAKER_(\d+)\s+\[\d{2}:\d{2}->\d{2}:\d{2}\]\s+SPEAKER_\d+\s+\|\s+",
        r"[\g<1>] ",
        transcript,
        flags=_re.MULTILINE,
    )
    system = (
        "You are a meeting secretary. Extract protocol from transcript ONLY from explicit statements. "
        "Return VALID JSON ONLY. NO markdown fences. NO trailing commas. Use DOUBLE QUOTES. "
        "Keys exactly: participants, agenda, decisions, assignments, open_questions, unclear. "
        "participants: array of {\"name\": \"SPEAKER_NN\", \"source_quote\": \"first line spoken by this speaker\"}. "
        "decisions: array of {\"text\": \"decision\", \"source_quote\": \"exact words from transcript\", \"approved_by\": [\"SPEAKER_NN\"]}. "
        "assignments: array of {\"task\": \"task\", \"assignee\": \"SPEAKER_NN\", \"deadline\": \"date or not_set\", \"source_quote\": \"exact words\"}. "
        "CRITICAL: name and assignee fields MUST contain ONLY SPEAKER_NN (SPEAKER_00, SPEAKER_01, etc). Never use real names. "
        "source_quote MUST be exact text from transcript — copy VERBATIM, do not paraphrase."
    )
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": clean_transcript},
            ],
            temperature=0.1,
        )
    except Exception as exc:
        _handle_exception(exc)
    content = resp.choices[0].message.content.strip()
    if content.startswith("```"):
        content = "\n".join(content.splitlines()[1:])
    if content.endswith("```"):
        content = "\n".join(content.splitlines()[:-1])
    try:
        return json.loads(content)
    except json.JSONDecodeError as exc:
        repaired = _repair_json(content)
        if repaired is not None:
            log.warning("LLM JSON repaired: %s", exc)
            return repaired
        fail(f"LLM returned invalid JSON for protocol: {exc}\nRaw: {content[:500]}")


def build_protocol(transcript: str, model: str, allow_cloud: bool) -> dict:
    if not _needs_protocol_chunking(transcript):
        return _build_protocol_chunk(transcript, model, allow_cloud)

    chunk_size_chars = _protocol_chunk_size_tokens() * PROTOCOL_CHARS_PER_TOKEN
    chunks = _split_protocol_transcript(transcript, chunk_size_chars)
    if len(chunks) == 1:
        return _build_protocol_chunk(transcript, model, allow_cloud)
    log.info("Building protocol from %s overlapping transcript chunks", len(chunks))
    return _merge_protocol_chunks(
        _build_protocol_chunk(chunk, model, allow_cloud) for chunk in chunks
    )


def _verify_protocol(
    protocol: dict, transcript: str, model: str, allow_cloud: bool
) -> dict:
    """Second-pass verification. Falls back to original protocol on failure."""
    if not _protocol_verification_enabled():
        return protocol
    enforce_cloud_policy(allow_cloud)
    from openai import OpenAI

    verify_url = os.getenv("MEETING_VERIFY_BASE_URL", LLM_BASE_URL)
    verify_key = os.getenv("MEETING_VERIFY_API_KEY", LLM_API_KEY)
    verify_model = os.getenv("MEETING_VERIFY_MODEL", model)
    # Use first 3000 chars of transcript as summary to avoid context overflow
    transcript_summary = transcript[:3000]
    prompt = (
        "Verify this protocol against the transcript excerpt. "
        "Find: missed decisions, wrong assignees, hallucinated participants. "
        "Output corrected JSON.\n\n"
        f"Protocol:\n{json.dumps(protocol, ensure_ascii=False)}\n\n"
        f"Transcript (first 3000 chars):\n{transcript_summary}"
    )
    client = OpenAI(base_url=verify_url, api_key=verify_key)
    try:
        content = (
            client.chat.completions.create(
                model=verify_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                timeout=60,
            )
            .choices[0]
            .message.content.strip()
        )
    except Exception as exc:
        log.warning("Protocol verification failed, using original: %s", exc)
        return protocol
    verified = _repair_json(content)
    if not isinstance(verified, dict):
        log.warning("Verification returned invalid JSON, using original")
        return protocol
    return verified


def _protocol_verification_enabled() -> bool:
    return os.getenv("MEETING_PROTOCOL_VERIFY", "true").lower() != "false"


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip().lower()


def _participant_map(value: Optional[str]) -> dict[str, str]:
    if not value:
        return {}
    return {
        speaker.strip(): name.strip()
        for pair in value.split(",")
        for speaker, name in [pair.strip().split("=", 1)]
    }


def replace_participant_labels(protocol: dict, participants: Optional[str]) -> None:
    participant_map = _participant_map(participants)
    for section in ["participants", "decisions", "assignments"]:
        for item in protocol.get(section, []):
            if not isinstance(item, dict):
                continue
            for field in ["name", "assignee", "approved_by"]:
                value = item.get(field)
                if isinstance(value, str) and value in participant_map:
                    item[field] = participant_map[value]
                elif isinstance(value, list):
                    item[field] = [participant_map.get(entry, entry) for entry in value]


def validate_protocol(protocol: Optional[dict], transcript: str) -> dict:
    if not protocol:
        return {
            "valid": False,
            "errors": ["protocol is empty"],
            "warnings": [],
            "overall_confidence": 0,
        }
    errors: List[str] = []
    warnings: List[str] = []
    transcript_norm = _normalize(transcript)
    for section in ["assignments", "decisions", "participants"]:
        for item in protocol.get(section, []):
            if isinstance(item, str):
                errors.append(f"{section} item is plain string, expected dict: {item[:80]}")
                continue
            sq = (item.get("source_quote") or "").strip()
            if not sq:
                errors.append(f"{section} item missing source_quote: {str(item)[:80]}")
                continue
            sq_norm = _normalize(sq)
            # Fuzzy match: quote words must appear in transcript
            sq_words = [w for w in sq_norm.split() if len(w) > 2]
            if sq_words:
                found = sum(1 for w in sq_words if w in transcript_norm)
                if found < max(1, len(sq_words) * 0.6):
                    errors.append(f"{section} source_quote not found: {sq[:80]}")
            assignee = (item.get("assignee") or "").strip()
            if assignee and assignee != "unknown":
                if assignee.lower() not in transcript_norm:
                    warnings.append(f"{section} assignee may be transliterated: {assignee}")
            deadline = (item.get("deadline") or "").strip()
            if deadline and deadline != "not_set":
                if not re.search(
                    r"(monday|tuesday|wednesday|thursday|friday|saturday|sunday|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|\d{1,2}[./:]\d{1,2}[./:]\d{2,4}|tmr|tomorrow|week|month|q[1-4])",
                    deadline,
                    re.I,
                ):
                    warnings.append(f"{section} deadline may be fabricated: {deadline}")
    confidence = 90
    if errors:
        confidence = min(confidence, 25)
    elif warnings:
        confidence = min(confidence, 70)
    return {
        "valid": len(errors) == 0,
        "errors": errors[:20],
        "warnings": warnings[:10],
        "overall_confidence": confidence,
    }


def atomic_write_json(path: Path, data: dict) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)
    return path


def write_text_docx(path: Path, title: str, paragraphs: Iterable[str]) -> None:
    """Write a small DOCX document from a title and plain-text paragraphs."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)
    for paragraph in paragraphs:
        text = str(paragraph).strip()
        if text:
            doc.add_paragraph(text)
    doc.save(path)
    log.info("Saved DOCX: %s", path)


def _docx_text(value: Any) -> str:
    if isinstance(value, dict):
        return value.get("text") or value.get("name") or json.dumps(
            value, ensure_ascii=False
        )
    return str(value)


def write_summary_docx(
    title: str,
    speaker: str,
    duration: str,
    topics: Iterable[Any],
    path: Path,
    key_concepts: Optional[Iterable[Any]] = None,
) -> None:
    """Write a meeting or lecture summary with topics and timestamped concepts."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)
    metadata = "; ".join(
        value for value in (f"Speaker: {speaker}" if speaker else "", f"Duration: {duration}" if duration else "") if value
    )
    if metadata:
        doc.add_paragraph(metadata)
    doc.add_heading("Key topics", level=1)
    for topic in topics:
        doc.add_paragraph(_docx_text(topic), style="List Bullet")
    concepts = list(key_concepts or [])
    if concepts:
        doc.add_heading("Key concepts", level=1)
        for concept in concepts:
            if isinstance(concept, dict):
                timestamp = concept.get("timestamp") or concept.get("time")
                text = concept.get("text") or concept.get("concept") or _docx_text(concept)
                doc.add_paragraph(
                    f"{timestamp}: {text}" if timestamp else text,
                    style="List Bullet",
                )
            else:
                doc.add_paragraph(str(concept), style="List Bullet")
    doc.save(path)
    log.info("Saved DOCX: %s", path)


def write_analytical_docx(sections: dict[str, str], path: Path) -> None:
    """Write an analytical note with the supplied section text."""
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Analytical note", level=0)
    for title, content in sections.items():
        doc.add_heading(str(title), level=1)
        for paragraph in str(content).split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    doc.save(path)
    log.info("Saved DOCX: %s", path)


def write_protocol_docx(protocol: dict, path: Path) -> None:
    if protocol.get("quality", {}).get("valid") is False:
        return
    from docx import Document
    from docx.shared import Pt, Cm

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    doc.add_heading("Протокол совещания", level=0)
    doc.add_paragraph("")

    if protocol.get("participants"):
        doc.add_heading("Участники", level=1)
        table = doc.add_table(rows=1 + len(protocol["participants"]), cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "№"
        table.rows[0].cells[1].text = "Участник"
        for idx, p in enumerate(protocol["participants"], 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = p.get("name", str(p))

    for section_name in ["agenda", "decisions", "assignments", "open_questions"]:
        items = protocol.get(section_name, [])
        if not items:
            continue
        doc.add_heading(section_name.replace("_", " ").title(), level=1)
        for idx, item in enumerate(items, 1):
            text = (
                item.get("text")
                or item.get("task")
                or item.get("question")
                or str(item)
            )
            doc.add_paragraph(f"{idx}. {text}")
    doc.save(path)
    log.info("Saved DOCX: %s", path)


def _now_iso() -> str:
    from datetime import datetime, timezone

    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def cmd_transcribe(args: argparse.Namespace) -> int:
    src = _resolve_source(args.source)
    if not src.exists():
        fail(f"File not found: {src}")
    check_resource_limits(src)
    audio = (
        src
        if src.suffix.lower() in {".wav", ".mp3", ".m4a", ".flac"}
        else src.with_suffix(".wav")
    )
    if audio != src:
        extract_audio(src, audio)
    transcript, meta = transcribe_audio(
        audio, args.model, args.language, args.device, args.compute_type
    )
    transcript = _clean_whisper_artifacts(transcript)
    out = Path(args.output) if args.output else src.with_suffix(".transcript.txt")
    out.write_text(transcript, encoding="utf-8")
    atomic_write_json(
        out.with_suffix(".transcript.json"), {"source_hash": sha256(src), **meta}
    )
    log.info("Saved transcript: %s", out)
    if _agent_mode_enabled():
        print(json.dumps(prepare_agent_transcript(transcript, out)))
    return 0


def cmd_translate(args: argparse.Namespace) -> int:
    src = Path(args.transcript)
    if not src.exists():
        fail(f"Transcript not found: {src}")
    lines = src.read_text(encoding="utf-8").splitlines()
    translated = translate_lines(lines, args.target_lang, allow_cloud=args.allow_cloud)
    out = Path(args.output) if args.output else src.with_suffix(".translated.txt")
    out.write_text("\n".join(translated), encoding="utf-8")
    log.info("Saved translation: %s", out)
    return 0


def cmd_protocol(args: argparse.Namespace) -> int:
    """Build a protocol with the legacy pipeline (legacy — prefer agent-driven via SKILL.md)."""
    src = Path(args.transcript)
    if not src.exists():
        fail(f"Transcript not found: {src}")
    transcript = src.read_text(encoding="utf-8")
    if _needs_protocol_chunking(transcript):
        log.info("Transcript exceeds 6000 tokens; protocol will be chunked")
    protocol = build_protocol(transcript, args.model, allow_cloud=args.allow_cloud)
    if _protocol_verification_enabled():
        protocol = _verify_protocol(
            protocol, transcript, args.model, allow_cloud=args.allow_cloud
        )
    validation = validate_protocol(protocol, transcript)
    replace_participant_labels(protocol, getattr(args, "participants", None))
    protocol["schema_version"] = "0.1.0"
    protocol["source_hash"] = sha256(src)
    protocol["stt_model"] = args.model
    protocol["llm_model"] = LLM_MODEL
    protocol["created_at"] = _now_iso()
    protocol["cloud_allowed"] = args.allow_cloud
    protocol["parameters"] = {
        "model": args.model,
        "allow_cloud": args.allow_cloud,
        "target_lang": getattr(args, "target_lang", None),
    }
    protocol["quality"] = validation
    out_path = Path(args.output) if args.output else src.with_suffix(".protocol.json")
    if validation["valid"]:
        atomic_write_json(out_path, protocol)
        if getattr(args, "docx", False):
            try:
                write_protocol_docx(protocol, src.with_suffix(".protocol.docx"))
            except PermissionError:
                from datetime import datetime
                fallback = src.with_suffix(f".protocol.{datetime.now().strftime('%H%M%S')}.docx")
                write_protocol_docx(protocol, fallback)
                log.warning("DOCX locked, saved to: %s", fallback)
        log.info("Saved protocol: %s", out_path)
        return 0
    rejected = out_path.with_suffix(".protocol.rejected.json")
    atomic_write_json(rejected, protocol)
    log.error("Invalid protocol saved to: %s", rejected)
    return 3


def cmd_process(args: argparse.Namespace) -> int:
    """Run the legacy end-to-end pipeline (legacy — prefer agent-driven via SKILL.md)."""
    src = _resolve_source(args.source)
    if not src.exists():
        fail(f"File not found: {src}")
    check_resource_limits(src)
    audio = (
        src
        if src.suffix.lower() in {".wav", ".mp3", ".m4a", ".flac"}
        else src.with_suffix(".wav")
    )
    if audio != src:
        extract_audio(src, audio)
    transcript, transcript_meta = transcribe_audio(
        audio, args.stt_model, args.language, args.device, args.compute_type
    )
    transcript = _clean_whisper_artifacts(transcript)
    transcript_path = src.with_suffix(".transcript.txt")
    transcript_path.write_text(transcript, encoding="utf-8")
    atomic_write_json(
        transcript_path.with_suffix(".transcript.json"),
        {"source_hash": sha256(src), **transcript_meta},
    )
    log.info("Saved transcript: %s", transcript_path)

    if not args.skip_translate:
        translated = translate_lines(
            transcript.splitlines(), args.target_lang, allow_cloud=args.allow_cloud
        )
        translated_path = src.with_suffix(".translated.txt")
        translated_path.write_text("\n".join(translated), encoding="utf-8")
        log.info("Saved translation: %s", translated_path)

    protocol = build_protocol(transcript, args.llm_model, allow_cloud=args.allow_cloud)
    validation = validate_protocol(protocol, transcript)
    replace_participant_labels(protocol, getattr(args, "participants", None))
    protocol["quality"] = validation
    protocol["schema_version"] = "0.1.0"
    protocol["source_hash"] = sha256(transcript_path)
    protocol["stt_model"] = args.stt_model
    protocol["llm_model"] = args.llm_model
    protocol["created_at"] = _now_iso()
    protocol["cloud_allowed"] = args.allow_cloud
    protocol["parameters"] = {
        "stt_model": args.stt_model,
        "llm_model": args.llm_model,
        "allow_cloud": args.allow_cloud,
        "target_lang": args.target_lang,
    }
    protocol_path = src.with_suffix(".protocol.json")
    if validation["valid"]:
        atomic_write_json(protocol_path, protocol)
        if args.docx:
            write_protocol_docx(protocol, src.with_suffix(".protocol.docx"))
        log.info("Saved protocol: %s", protocol_path)
    else:
        rejected = protocol_path.with_suffix(".protocol.rejected.json")
        atomic_write_json(rejected, protocol)
        log.error("Invalid protocol saved to: %s", rejected)
    if not validation["valid"]:
        return 3
    return 0


def main() -> int:
    p = argparse.ArgumentParser(description="Meeting Intelligence CLI")
    sub = p.add_subparsers(dest="command")

    transcribe_p = sub.add_parser("transcribe")
    transcribe_p.add_argument("source")
    transcribe_p.add_argument("--model", default=TRANSCRIBE_MODEL)
    transcribe_p.add_argument("--language", default=TRANSCRIBE_LANG)
    transcribe_p.add_argument("--device", default=TRANSCRIBE_DEVICE)
    transcribe_p.add_argument("--compute-type", default=TRANSCRIBE_COMPUTE)
    transcribe_p.add_argument("--output", type=Path, default=None)

    translate_p = sub.add_parser("translate")
    translate_p.add_argument("transcript", type=Path)
    translate_p.add_argument("--target-lang", default="ru")
    translate_p.add_argument("--allow-cloud", action="store_true", default=False)
    translate_p.add_argument("--output", type=Path, default=None)

    agent_transcript_p = sub.add_parser(
        "agent-transcript",
        help="Emit a cleaned transcript JSON payload for agent consumption",
    )
    agent_transcript_p.add_argument("transcript", type=Path)
    agent_transcript_p.add_argument("--docx", action="store_true", default=False)
    agent_transcript_p.add_argument("--output", type=Path, default=None)

    generate_docx_p = sub.add_parser(
        "generate-docx", help="Generate DOCX from summary, analytical, or protocol content"
    )
    generate_docx_p.add_argument("--type", choices=("summary", "analytical", "protocol"), required=True)
    generate_docx_p.add_argument("--input", type=Path, required=True)
    generate_docx_p.add_argument("--output", type=Path, required=True)

    protocol_p = sub.add_parser("protocol")
    protocol_p.add_argument("transcript", type=Path)
    protocol_p.add_argument("--model", default=LLM_MODEL)
    protocol_p.add_argument("--allow-cloud", action="store_true", default=False)
    protocol_p.add_argument("--docx", action="store_true", default=False)
    protocol_p.add_argument("--participants", default=None, help="SPEAKER_00=Имя,SPEAKER_01=Имя,...")
    protocol_p.add_argument("--output", type=Path, default=None)

    process_p = sub.add_parser("process")
    process_p.add_argument("source")
    process_p.add_argument("--stt-model", default=TRANSCRIBE_MODEL)
    process_p.add_argument("--llm-model", default=LLM_MODEL)
    process_p.add_argument("--language", default=TRANSCRIBE_LANG)
    process_p.add_argument("--device", default=TRANSCRIBE_DEVICE)
    process_p.add_argument("--compute-type", default=TRANSCRIBE_COMPUTE)
    process_p.add_argument("--target-lang", default="ru")
    process_p.add_argument("--skip-translate", action="store_true", default=False)
    process_p.add_argument("--docx", action="store_true", default=False)
    process_p.add_argument("--allow-cloud", action="store_true", default=False)
    process_p.add_argument("--participants", default=None, help="SPEAKER_00=Name,SPEAKER_01=Name,...")
    process_p.add_argument("--output", type=Path, default=None)

    args = p.parse_args()
    if not args.command:
        p.print_help()
        return 2

    try:
        if args.command == "transcribe":
            return cmd_transcribe(args)
        if args.command == "translate":
            return cmd_translate(args)
        if args.command == "agent-transcript":
            return cmd_agent_transcript(args)
        if args.command == "generate-docx":
            return cmd_generate_docx(args)
        if args.command == "protocol":
            return cmd_protocol(args)
        if args.command == "process":
            return cmd_process(args)
    except MeetingError as exc:
        fail(str(exc))
    fail("Unknown command")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
