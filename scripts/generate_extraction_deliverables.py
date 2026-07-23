"""Create evidence-linked meeting and lecture deliverables from the four transcripts."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, PatternFill


ROOT = Path(r"C:\Work\Assist\meeting")
OUT = ROOT / "pipeline"
LINE = re.compile(r"^\[(?P<time>[^]]+)\]\s+(?P<speaker>[^|]+)\|\s*(?P<text>.*)$")


def source(prefix: str) -> Path:
    return next(ROOT.glob(prefix + "*.transcript.txt"))


ITEMS = [
    {
        "path": source("1 лекция"),
        "kind": "lecture",
        "language": "English",
        "title": "Lecture notes: AI knowledge bases and legal work",
        "speaker": "Multiple labeled speakers (SPEAKER_00 to SPEAKER_15); names are not stated.",
        "topics": [
            (
                "00:04",
                "Keeping a model's knowledge current when a document or requirement changes.",
            ),
            ("05:25", "Using AI to shorten and search material for a needed fact."),
            ("08:46", "Building a project knowledge base and retaining task context."),
            (
                "14:00",
                "Working with multiple document versions and loading documents into a tool.",
            ),
            (
                "20:50",
                "Prompting against a large body of knowledge and reducing routine legal work.",
            ),
            ("25:19", "Prompt examples and task-specific inputs."),
        ],
        "concepts": [
            (
                "Document-grounded answers",
                "A changed document makes new knowledge available to the model.",
                "[00:09->00:24] SPEAKER_00 | So, you have changed the document... this knowledge is available to the model... and generate answers with the new knowledge.",
            ),
            (
                "Knowledge base",
                "The lecture connects a project knowledge base to finding information and carrying task context.",
                "[08:46->09:44] SPEAKER_03 | where we talked about the base of the knowledge... our base of knowledge of this project",
            ),
            (
                "Document versions",
                "AI work must handle the working document and its versions, not only a generic answer.",
                "[13:50->14:21] SPEAKER_04 | ...answer to a question. After that, we need to work with a document. We have several versions, for example,",
            ),
        ],
        "examples": [
            (
                "Searching a concise record",
                "[05:38->05:46] SPEAKER_03 | conspect... in a short way to find some information which was needed.",
            ),
            (
                "Prompt design",
                "[25:19->26:08] SPEAKER_05 | For example, a prompt... The next example of a prompt",
            ),
        ],
        "qa": "A question-oriented exchange appears near 27:56, but the automatic transcript does not preserve a reliable question and answer pair.",
    },
    {
        "path": source("2 лекция"),
        "kind": "lecture",
        "language": "English",
        "title": "Lecture notes: Legal AI workflows and answer reliability",
        "speaker": "Multiple labeled speakers (SPEAKER_00 to SPEAKER_05); names are not stated.",
        "topics": [
            (
                "03:01",
                "Legal questions can have several possible outcomes and uncertainty.",
            ),
            ("05:31", "A response is not guaranteed to be completely correct."),
            ("07:04", "A multi-step model process for constructing an answer."),
            (
                "33:20",
                "Comparing information in a single table and checking simple questions.",
            ),
            ("37:01", "Using documents, law, and agreements to answer a question."),
            ("38:35", "Legal conflict and risk as evaluation points."),
        ],
        "concepts": [
            (
                "Uncertainty",
                "The lecture explicitly treats legal answers as less than fully certain.",
                "[05:31->05:37] SPEAKER_01 | I think that is for all of us, which is not a 100% correct answer,",
            ),
            (
                "Multi-step answering",
                "The answer process is described as a sequence of questions and answers rather than one direct output.",
                "[07:04->07:17] SPEAKER_02 | We have a model... a more difficult process of the answer. First, we are going to be a question for the answer",
            ),
            (
                "Source documents",
                "The model should combine material from law and an agreement when answering.",
                "[37:14->37:25] SPEAKER_04 | He will take it from the law and from the agreement... or a big document",
            ),
        ],
        "examples": [
            (
                "Document-based legal answer",
                "[37:01->37:09] SPEAKER_04 | by using of the document, which can be used to answer",
            )
        ],
        "qa": "The transcript contains question phrasing, but no reliably transcribed standalone Q&A sequence.",
    },
    {
        "path": source("15330497137360"),
        "kind": "meeting",
        "language": "Русский",
        "title": "Рабочая встреча: замечания к методикам и показателям",
        "participants": "Андрей Васильевич, Алексей Николаевич Яковлев, Петр; остальные участники обозначены как SPEAKER_00.",
        "agenda": [
            (
                "Рассмотреть замечания к методикам и показателям.",
                "[00:50->01:06] SPEAKER_00 | Давайте пройдем по замечаниям... 130 страниц.",
            ),
            (
                "Уточнить применимость ГОСТов, формул, источников данных и порядка верификации.",
                "[02:49->03:19] SPEAKER_00 | мы ссылаемся на какие-то там гостые методики... уточните, что там не надо применять весь гост...",
            ),
            (
                "Определить порядок направления доработанного варианта на согласование.",
                "[24:21->25:06] SPEAKER_00 | пришлете в Минюст... исправный вариант... руководство само ответит... либо мы согласуем сами.",
            ),
        ],
        "decisions": [
            (
                "Для вопросов, не относящихся к методике расчета, дать пояснение и снять замечание.",
                "[10:26->10:57] SPEAKER_00 | ...в примечании напишете, что состав показателей определяется Минцифрой как держателем данных... Просто объяснение такое... и все, мы это все снимаем. Зафиксировал, спасибо.",
            ),
            (
                "В методике отразить деление на нуль и отсутствие данных.",
                "[12:40->12:47] SPEAKER_00 | деление на нуль и отсутствие данных мы в методике отразим",
            ),
            (
                "Направлять в Минюст исправленный вариант, а не детальный ответ на каждое замечание.",
                "[24:21->24:55] SPEAKER_00 | пришлете в Минюст не нам уже исправный вариант... не надо детального ответа, нужно скорректированный... вариант",
            ),
        ],
        "assignments": [
            (
                "Уточнить, какие части ГОСТа применяются при подготовке исправленных вариантов.",
                "не указан",
                "не указан",
                "[03:01->03:19] SPEAKER_00 | пожалуйста, уточните, что там не надо применять весь гост... а только применяем в части 3, 4, 5.",
            ),
            (
                "Доработать неуточненные моменты и направить доработанную версию на согласование.",
                "не указан",
                "не указан",
                "[00:33->00:50] SPEAKER_00 | с нашей стороны доработаем те моменты... и будем направлять... доработанную версию на согласование.",
            ),
        ],
        "open": [
            (
                "Насколько детально описывать сбор данных и уровень детализации.",
                "[08:36->09:10] SPEAKER_00 | этот вопрос сейчас... Насколько здесь принципиально детализация, на какой уровень",
            ),
            (
                "Как классифицировать программные решения и исключения расчета.",
                "[21:20->22:26] SPEAKER_00 | не разработал критерии... вопросы классификации, программ... будут уточняться",
            ),
        ],
        "risks": [
            (
                "Технические ошибки в формулах, процентах и долях.",
                "[01:39->02:07] SPEAKER_00 | много небрежности технических... путаете проценты и доли",
            ),
            (
                "Неприменимость требований ГОСТа без уточнения их части.",
                "[02:49->03:13] SPEAKER_00 | ...может быть и неприменимо... не надо применять весь гост",
            ),
        ],
        "next": [
            (
                "Подготовить скорректированный итоговый вариант и направить его в Минюст.",
                "[24:21->25:06] SPEAKER_00 | пришлете в Минюст... исправный вариант",
            )
        ],
    },
    {
        "path": source("2026_04_09"),
        "kind": "meeting",
        "language": "Русский",
        "title": "Рабочая встреча: KPI внедрения ИИ в спорте",
        "participants": "Ильдар Ахмедов, Дмитрий Григоренко, Максут Шадаев, Олег Игоревич, Александр Иванович, Надежда Викторовна; остальные участники обозначены как SPEAKER_00–SPEAKER_04.",
        "agenda": [
            (
                "Сформировать отраслевые KPI внедрения ИИ в спорте на 2027–2030 годы.",
                "[00:30->01:02] SPEAKER_00 | был дан старт формирования ключевых показателей эффективности по внедрению искусственной интеллекта на 27-30 годы.",
            ),
            (
                "Обсудить стартовые предложения по каждому показателю.",
                "[05:49->06:29] SPEAKER_00 | давайте мы пройдем по каждой строке... Принимается этот показатель или не принимается",
            ),
            (
                "Собрать предложения для следующей итерации и подготовки к стратегической сессии.",
                "[01:19->03:09] SPEAKER_00 | На следующей неделе мы с вами покажем показатели... на неделе с 24 апреля проводятся большие стратегические сессии",
            ),
        ],
        "decisions": [
            (
                "Взять все высказанные предложения в работу для следующей встречи.",
                "[111:37->111:45] SPEAKER_04 | тогда, берем все вместе в работу и до встречи на следующей неделе.",
            ),
            (
                "На следующей встрече рассматривать каждый показатель отдельным листом с критериями, эффектами и значениями.",
                "[110:46->111:02] SPEAKER_04 | задача сейчас их... детализировать, написать по нему критерию и на следующей встрече посмотреть уже один лист, один показатель... со всеми эффектами, значениями",
            ),
        ],
        "assignments": [
            (
                "Направить предложения по формулировкам, показателям и направлениям в машиночитаемом виде для следующей итерации.",
                "не указан",
                "завтра до 11–12 часов",
                "[111:47->112:10] SPEAKER_04 | просим... завтра нам часов до 12 до 11 направить... в машиночтайном виде нам эти предложения сброситься",
            ),
            (
                "Направить сформулированные пункты перед следующей встречей экспертам по ИИ в спорте для комментариев.",
                "не указан",
                "до следующей встречи",
                "[110:14->110:30] SPEAKER_04 | перед следующей встречей... пункты... прислать... готовы быстро дать комментарии",
            ),
        ],
        "open": [
            (
                "Какие из стартовых показателей принять, доуточнить или заменить.",
                "[05:55->06:11] SPEAKER_00 | Принимается этот показатель или не принимается... может быть доформулирован, доуточнен",
            ),
            (
                "Как сформулировать показатель выявления талантов без риска дискриминации.",
                "[96:12->98:11] SPEAKER_04 | Давайте очень аккуратно это слово употреблять... с условием дискриминации, чтобы никакого составляющей не было",
            ),
        ],
        "risks": [
            (
                "Немашиночитаемые документы и ручная проверка создают ошибки.",
                "[21:36->21:58] SPEAKER_01 | протокола соревнований... в немашинно-читаемом виде, и там идет ручная проверка документов, и там зачастую возникают ошибки",
            ),
            (
                "Для KPI могут потребоваться изменения нормативного регулирования, данные и инфраструктура.",
                "[04:11->04:35] SPEAKER_00 | Нужно, например, менять нормативку... необходимы... данные... инфраструктура, датчики, камеры",
            ),
        ],
        "next": [
            (
                "Детализировать KPI, критерии, эффекты и значения к следующей встрече.",
                "[110:46->111:02] SPEAKER_04 | задача сейчас их... детализировать",
            ),
            (
                "Провести следующую рабочую встречу на следующей неделе.",
                "[111:37->111:45] SPEAKER_04 | до встречи на следующей неделе",
            ),
        ],
    },
]


def make_doc(title: str) -> Document:
    d = Document()
    d.styles["Normal"].font.name = "Arial"
    d.styles["Normal"].font.size = Pt(10)
    d.add_heading(title, 0)
    return d


def quote(d: Document, text: str) -> None:
    p = d.add_paragraph("Источник: " + text)
    p.style = "Intense Quote"


def add_items(d: Document, heading: str, values: list[tuple]) -> None:
    d.add_heading(heading, 1)
    for value in values:
        d.add_paragraph(value[0], style="List Bullet")
        quote(d, value[-1])


def transcript_lines(path: Path) -> list[str]:
    return [x for x in path.read_text(encoding="utf-8-sig").splitlines() if x.strip()]


def write_minutes(item: dict, folder: Path) -> None:
    d = make_doc(
        "Подробный протокол" if item["kind"] == "meeting" else "Detailed minutes"
    )
    d.add_paragraph("Source transcript: " + item["path"].name)
    for raw in transcript_lines(item["path"]):
        d.add_paragraph(raw)
    d.save(folder / "detailed-minutes.docx")


def write_lecture(item: dict, folder: Path) -> None:
    def topic_doc(name: str, brief: bool = False) -> Document:
        d = make_doc(name)
        d.add_paragraph("Language: English")
        d.add_paragraph("Speaker: " + item["speaker"])
        d.add_heading("Key topics with timestamps", 1)
        for t, text in item["topics"]:
            d.add_paragraph(f"[{t}] {text}", style="List Bullet")
        return d

    d = topic_doc("Summary")
    d.add_heading("Key concepts", 1)
    for name, text, ev in item["concepts"]:
        d.add_paragraph(f"{name}. {text}", style="List Bullet")
        quote(d, ev)
    d.add_heading("Q&A", 1)
    d.add_paragraph(item["qa"])
    d.save(folder / "summary.docx")
    d = topic_doc("Analytical note")
    d.add_heading("Analysis", 1)
    d.add_paragraph(
        "The lecture presents document-grounded AI as practical support for legal work. The transcript is recognition-degraded, so speaker attribution and exact wording should be checked against the recording before use as a formal source."
    )
    d.add_heading("Examples", 1)
    for n, ev in item["examples"]:
        d.add_paragraph(n, style="List Bullet")
        quote(d, ev)
    d.save(folder / "analytical.docx")
    d = topic_doc("Knowledge article: legal AI workflows")
    d.add_heading("Practical takeaway", 1)
    d.add_paragraph(
        "Keep the relevant document set and version context available when asking a model for a work answer. Treat responses as reviewable working material, especially in legal matters."
    )
    d.save(folder / "knowledge-article.docx")
    d = topic_doc("Executive brief")
    d.add_paragraph(
        "The recording focuses on AI supported by document knowledge, prompt design, and reliability checks in legal work. The material supports a document-first workflow, not unsupervised reliance on a generic answer."
    )
    d.save(folder / "executive-brief.docx")


def write_sheet(path: Path, headers: list[str], rows: list[tuple]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for c in ws[1]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="1F4E78")
    for row in rows:
        ws.append(row)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = min(
            85, max(14, max(len(str(c.value or "")) for c in col) + 2)
        )
    wb.save(path)


def write_meeting(item: dict, folder: Path) -> None:
    d = make_doc("Протокол")
    d.add_paragraph("Язык: русский")
    d.add_paragraph("Участники: " + item["participants"])
    for h, v in [
        ("Повестка", item["agenda"]),
        ("Решения", item["decisions"]),
        ("Поручения", item["assignments"]),
        ("Открытые вопросы", item["open"]),
        ("Риски", item["risks"]),
        ("Следующие шаги", item["next"]),
    ]:
        add_items(d, h, v)
    d.save(folder / "protocol.docx")
    write_sheet(
        folder / "decision-register.xlsx",
        ["Решение", "Источник (дословная цитата)"],
        item["decisions"],
    )
    write_sheet(
        folder / "assignment-list.xlsx",
        ["Поручение", "Исполнитель", "Срок", "Источник (дословная цитата)"],
        item["assignments"],
    )
    d = make_doc("Саммари")
    d.add_paragraph("Участники: " + item["participants"])
    add_items(d, "Главные темы", item["agenda"])
    add_items(d, "Согласованные результаты", item["decisions"])
    d.save(folder / "summary.docx")
    d = make_doc("Аналитическая записка")
    add_items(d, "Риски и ограничения", item["risks"])
    add_items(d, "Вопросы, требующие проработки", item["open"])
    d.add_heading("Вывод", 1)
    d.add_paragraph(
        "Материал фиксирует рабочее обсуждение. Решения и поручения внесены только там, где запись содержит явное согласование или прямой запрос."
    )
    d.save(folder / "analytical.docx")
    d = make_doc("План действий")
    add_items(d, "Действия", item["assignments"])
    add_items(d, "Ближайшие шаги", item["next"])
    d.save(folder / "action-plan.docx")
    d = make_doc("Справка для руководства")
    d.add_paragraph("Тема: " + item["title"])
    add_items(d, "Согласованные действия", item["decisions"])
    add_items(d, "Критичные риски", item["risks"])
    d.save(folder / "executive-brief.docx")


def validate(folder: Path, required: list[str]) -> None:
    for name in required:
        p = folder / name
        if not p.exists() or p.stat().st_size == 0:
            raise RuntimeError(f"Missing output: {p}")
        if p.suffix == ".docx":
            Document(p)
        if p.suffix == ".xlsx":
            load_workbook(p, read_only=True).close()


def main() -> None:
    for item in ITEMS:
        folder = OUT / item["path"].stem.replace(".transcript", "")
        folder.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(item["path"], folder / "transcript.txt")
        write_minutes(item, folder)
        if item["kind"] == "lecture":
            write_lecture(item, folder)
            validate(
                folder,
                [
                    "summary.docx",
                    "analytical.docx",
                    "detailed-minutes.docx",
                    "knowledge-article.docx",
                    "executive-brief.docx",
                ],
            )
        else:
            write_meeting(item, folder)
            validate(
                folder,
                [
                    "protocol.docx",
                    "summary.docx",
                    "analytical.docx",
                    "decision-register.xlsx",
                    "assignment-list.xlsx",
                    "detailed-minutes.docx",
                    "action-plan.docx",
                    "executive-brief.docx",
                ],
            )


if __name__ == "__main__":
    main()
