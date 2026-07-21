#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import logging
import os
import subprocess
import sys
from pathlib import Path
from typing import Optional

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("meeting")

ALLOW_CLOUD = os.getenv("MEETING_ALLOW_CLOUD", "false").lower() == "true"
MAX_FILE_MB = int(os.getenv("MEETING_MAX_FILE_MB", "2048"))
MAX_DURATION_SEC = int(os.getenv("MEETING_MAX_DURATION_SEC", "7200"))
TRANSCRIBE_MODEL = os.getenv("MEETING_TRANSCRIBE_MODEL", "small")
TRANSCRIBE_DEVICE = os.getenv("MEETING_TRANSCRIBE_DEVICE", "cpu")
TRANSCRIBE_COMPUTE = os.getenv("MEETING_TRANSCRIBE_COMPUTE", "int8")
TRANSCRIBE_LANG = os.getenv("MEETING_TRANSCRIBE_LANG", "en")
LLM_BASE_URL = os.getenv("MEETING_LLM_BASE_URL", "http://localhost:1234/v1")
LLM_API_KEY = os.getenv("MEETING_LLM_API_KEY", "lm-studio")
LLM_MODEL = os.getenv("MEETING_LLM_MODEL", "qwen2.5-7b-instruct")


def fail(message: str, code: int = 2) -> None:
    log.error(message)
    raise SystemExit(code)


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def is_loopback_url(url: str) -> bool:
    """Разрешаем только loopback/localhost для конфиденциальных данных."""
    from urllib.parse import urlparse
    host = (urlparse(url).hostname or "").lower()
    return host in {"127.0.0.1", "localhost", "::1", ""}


def enforce_cloud_policy(allow_cloud: bool) -> None:
    if not allow_cloud:
        if not is_loopback_url(LLM_BASE_URL):
            fail(
                f"Cloud LLM is disabled (MEETING_ALLOW_CLOUD=false). "
                f"HOST={LLM_BASE_URL!r}. Pass --allow-cloud explicitly to enable."
            )


def check_resource_limits(path: Path) -> None:
    size_mb = path.stat().st_size / (1024 * 1024)
    if size_mb > MAX_FILE_MB:
        fail(f"File too large: {size_mb:.1f} MB > {MAX_FILE_MB} MB")

    if path.suffix.lower() in {".mp4", ".mov", ".webm", ".mkv"}:
        try:
            probe = subprocess.run(
                ["ffprobe", "-v", "quiet", "-show_format", str(path)],
                capture_output=True, text=True, timeout=30, check=False,
            )
            if probe.returncode != 0:
                log.warning("ffprobe failed: %s", probe.stderr[-200:])
            else:
                for line in probe.stdout.splitlines():
                    if line.startswith("duration="):
                        dur = float(line.split("=", 1)[1])
                        if dur > MAX_DURATION_SEC:
                            fail(f"Duration too long: {dur:.0f}s > {MAX_DURATION_SEC}s")
                        break
        except Exception as exc:
            log.warning("Duration check skipped: %s", exc)

    if path.suffix.lower() in {".wav", ".mp3", ".m4a", ".flac"}:
        try:
            probe = subprocess.run(
                ["ffprobe", "-v", "quiet", "-show_format", str(path)],
                capture_output=True, text=True, timeout=30, check=False,
            )
            if probe.returncode == 0:
                for line in probe.stdout.splitlines():
                    if line.startswith("duration="):
                        dur = float(line.split("=", 1)[1])
                        if dur > MAX_DURATION_SEC:
                            fail(f"Duration too long: {dur:.0f}s > {MAX_DURATION_SEC}s")
                        break
        except Exception as exc:
            log.warning("Duration check skipped: %s", exc)


def extract_audio(src: Path, dst: Path) -> Path:
    cmd = [
        "ffmpeg", "-y", "-i", str(src),
        "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1",
        "-t", str(MAX_DURATION_SEC), str(dst),
    ]
    log.info("Extracting audio -> %s", dst)
    res = subprocess.run(cmd, capture_output=True, text=True, timeout=600, check=False)
    if res.returncode != 0:
        fail(f"ffmpeg failed: {res.stderr[-400:]}")
    if not dst.exists() or dst.stat().st_size == 0:
        fail("ffmpeg produced empty audio file")
    return dst


def transcribe_audio(audio: Path, model: str, language: Optional[str], device: str, compute_type: str) -> str:
    from faster_whisper import WhisperModel
    log.info("Loading whisper model=%s device=%s compute_type=%s", model, device, compute_type)
    m = WhisperModel(model, device=device, compute_type=compute_type)
    segments, info = m.transcribe(
        str(audio),
        language=language,
        beam_size=5,
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500),
    )
    log.info("Detected language=%s duration=%.1fs", info.language, info.duration)
    out = []
    for seg in segments:
        ts = f"[{int(seg.start//60):02d}:{int(seg.start%60):02d}]"
        text = seg.text.strip()
        if text:
            out.append(f"{ts}|{text}")
    return "\n".join(out)


def translate_lines(lines, target_lang: str, *, allow_cloud: bool) -> list[str]:
    enforce_cloud_policy(allow_cloud)
    if not lines:
        return []
    try:
        from openai import OpenAI
    except Exception as exc:
        fail(f"openai package missing: {exc}")

    client = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)
    out = []
    failed = 0
    for line in lines:
        if "|" not in line:
            continue
        ts, text = line.split("|", 1)
        prompt = (
            f"Translate the following meeting transcript sentence into {target_lang}. "
            f"Keep names, codes, and technical terms unchanged. Output ONLY the translated sentence, no quotes, no extra text.\n\n{text}"
        )
        try:
            resp = client.chat.completions.create(
                model=LLM_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )
            out.append(f"{ts}|{resp.choices[0].message.content.strip()}")
        except Exception as exc:
            log.warning("translate failed for line: %s", exc)
            failed += 1
    if failed:
        fail(f"LLM translation failed for {failed} line(s)")
    return out


def build_protocol(transcript: str, model: str, *, allow_cloud: bool) -> dict:
    enforce_cloud_policy(allow_cloud)
    try:
        from openai import OpenAI
    except Exception as exc:
        fail(f"openai package missing: {exc}")

    client = OpenAI(base_url=LLM_BASE_URL, api_key=LLM_API_KEY)
    system = (
        "You are a meeting secretary. Extract protocol from transcript ONLY from explicit statements. "
        "Return JSON with keys: participants, agenda, decisions, assignments, open_questions, unclear. "
        "Every item must have source_quote. If assignee or deadline is missing, use unknown/not_set."
    )
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": transcript}],
        temperature=0.1,
    )
    content = resp.choices[0].message.content.strip()
    if content.startswith("```"):
        content = "\n".join(content.splitlines()[1:])
    if content.endswith("```"):
        content = "\n".join(content.splitlines()[:-1])
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        fail("LLM returned invalid JSON for protocol")


def validate_protocol(protocol: Optional[dict], transcript: str) -> dict:
    if not protocol:
        return {"valid": False, "errors": ["protocol is empty"], "warnings": [], "overall_confidence": 0}
    errors = []
    warnings = []
    transcript_lower = transcript.lower()
    for section in ["assignments", "decisions", "participants"]:
        for item in protocol.get(section, []):
            sq = item.get("source_quote", "").strip()
            if not sq:
                errors.append(f"{section} item missing source_quote: {str(item)[:80]}")
                continue
            if sq.lower() not in transcript_lower:
                words = [w for w in sq.lower().split() if len(w) > 3]
                if words:
                    ratio = sum(1 for w in words if w in transcript_lower) / len(words)
                    if ratio < 0.6:
                        errors.append(f"{section} source_quote not found: {sq[:80]}")
                    else:
                        warnings.append(f"{section} source_quote partial match: {sq[:80]}")
    confidence = 90 if not errors else (50 if not warnings else 70)
    return {
        "valid": len(errors) == 0,
        "errors": errors,
        "warnings": warnings[:10],
        "overall_confidence": confidence,
    }


def atomic_write_json(path: Path, data: dict) -> None:
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    tmp.replace(path)


def cmd_transcribe(args: argparse.Namespace) -> int:
    src = Path(args.source)
    if not src.exists():
        fail(f"File not found: {src}")
    check_resource_limits(src)
    audio = src if src.suffix.lower() in {".wav", ".mp3", ".m4a", ".flac"} else src.with_suffix(".wav")
    if audio != src:
        extract_audio(src, audio)
    text = transcribe_audio(audio, args.model, args.language, args.device, args.compute_type)
    out = Path(args.output) if args.output else src.with_suffix(".transcript.txt")
    out.write_text(text, encoding="utf-8")
    log.info("Saved transcript: %s", out)
    return 0


def cmd_translate(args: argparse.Namespace) -> int:
    src = Path(args.transcript)
    if not src.exists():
        fail(f"Transcript not found: {src}")
    lines = src.read_text(encoding="utf-8").splitlines()
    translated = translate_lines(lines, args.target_lang, allow_cloud=args.allow_cloud)
    out = Path(args.output) if args.output else src.with_suffix(".translated.txt")
    out.write_text("\n".join(translated), encoding="utf-8")
    log.info("Saved translation: %s", out)
    return 0


def cmd_protocol(args: argparse.Namespace) -> int:
    src = Path(args.transcript)
    if not src.exists():
        fail(f"Transcript not found: {src}")
    transcript = src.read_text(encoding="utf-8")
    protocol = build_protocol(transcript, args.model, allow_cloud=args.allow_cloud)
    validation = validate_protocol(protocol, transcript)
    protocol["quality"] = validation
    out_path = Path(args.output) if args.output else src.with_suffix(".protocol.json")
    if validation["valid"]:
        atomic_write_json(out_path, protocol)
        log.info("Saved protocol: %s", out_path)
        return 0
    rejected = out_path.with_suffix(".protocol.rejected.json")
    atomic_write_json(rejected, protocol)
    log.error("Invalid protocol saved to: %s", rejected)
    return 3


def cmd_process(args: argparse.Namespace) -> int:
    src = Path(args.source)
    if not src.exists():
        fail(f"File not found: {src}")
    check_resource_limits(src)
    audio = src if src.suffix.lower() in {".wav", ".mp3", ".m4a", ".flac"} else src.with_suffix(".wav")
    if audio != src:
        extract_audio(src, audio)
    transcript = transcribe_audio(audio, args.model, args.language, args.device, args.compute_type)
    transcript_path = src.with_suffix(".transcript.txt")
    transcript_path.write_text(transcript, encoding="utf-8")
    log.info("Saved transcript: %s", transcript_path)

    if not args.skip_translate:
        translated = translate_lines(transcript.splitlines(), args.target_lang, allow_cloud=args.allow_cloud)
        translated_path = src.with_suffix(".translated.txt")
        translated_path.write_text("\n".join(translated), encoding="utf-8")
        log.info("Saved translation: %s", translated_path)

    protocol = build_protocol(transcript, args.model, allow_cloud=args.allow_cloud)
    validation = validate_protocol(protocol, transcript)
    protocol["quality"] = validation
    protocol_path = src.with_suffix(".protocol.json")
    if validation["valid"]:
        atomic_write_json(protocol_path, protocol)
        log.info("Saved protocol: %s", protocol_path)
    else:
        rejected = protocol_path.with_suffix(".protocol.rejected.json")
        atomic_write_json(rejected, protocol)
        log.error("Invalid protocol saved to: %s", rejected)

    if args.docx:
        write_protocol_docx(protocol, src.with_suffix(".protocol.docx"))
        log.info("Saved docx: %s", src.with_suffix(".protocol.docx"))

    if not validation["valid"]:
        return 3
    return 0


def write_protocol_docx(protocol: dict, path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, Cm
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    doc.add_heading("Протокол совещания", level=0)
    doc.add_paragraph("")

    if protocol.get("participants"):
        doc.add_heading("Участники", level=1)
        table = doc.add_table(rows=1 + len(protocol["participants"]), cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "№"
        table.rows[0].cells[1].text = "Участник"
        for idx, p in enumerate(protocol["participants"], 1):
            table.rows[idx].cells[0].text = str(idx)
            table.rows[idx].cells[1].text = p.get("name", str(p))

    for section_name in ["agenda", "decisions", "assignments", "open_questions"]:
        items = protocol.get(section_name, [])
        if not items:
            continue
        doc.add_heading(section_name.replace("_", " ").title(), level=1)
        for idx, item in enumerate(items, 1):
            text = item.get("text") or item.get("task") or item.get("question") or str(item)
            doc.add_paragraph(f"{idx}. {text}")

    doc.save(path)


def main() -> int:
    p = argparse.ArgumentParser(description="Meeting Intelligence CLI")
    sub = p.add_subparsers(dest="command")

    transcribe_p = sub.add_parser("transcribe")
    transcribe_p.add_argument("source", type=Path)
    transcribe_p.add_argument("--model", default=TRANSCRIBE_MODEL)
    transcribe_p.add_argument("--language", default=TRANSCRIBE_LANG)
    transcribe_p.add_argument("--device", default=TRANSCRIBE_DEVICE)
    transcribe_p.add_argument("--compute-type", default=TRANSCRIBE_COMPUTE)
    transcribe_p.add_argument("--output", type=Path, default=None)

    translate_p = sub.add_parser("translate")
    translate_p.add_argument("transcript", type=Path)
    translate_p.add_argument("--target-lang", default="ru")
    translate_p.add_argument("--allow-cloud", action="store_true", default=False)
    translate_p.add_argument("--output", type=Path, default=None)

    protocol_p = sub.add_parser("protocol")
    protocol_p.add_argument("transcript", type=Path)
    protocol_p.add_argument("--model", default=LLM_MODEL)
    protocol_p.add_argument("--allow-cloud", action="store_true", default=False)
    protocol_p.add_argument("--output", type=Path, default=None)

    process_p = sub.add_parser("process")
    process_p.add_argument("source", type=Path)
    process_p.add_argument("--model", default=TRANSCRIBE_MODEL)
    process_p.add_argument("--language", default=TRANSCRIBE_LANG)
    process_p.add_argument("--device", default=TRANSCRIBE_DEVICE)
    process_p.add_argument("--compute-type", default=TRANSCRIBE_COMPUTE)
    process_p.add_argument("--target-lang", default="ru")
    process_p.add_argument("--skip-translate", action="store_true", default=False)
    process_p.add_argument("--docx", action="store_true", default=False)
    process_p.add_argument("--allow-cloud", action="store_true", default=False)
    process_p.add_argument("--output", type=Path, default=None)

    args = p.parse_args()
    if not args.command:
        p.print_help()
        return 2

    if args.command == "transcribe":
        return cmd_transcribe(args)
    if args.command == "translate":
        return cmd_translate(args)
    if args.command == "protocol":
        return cmd_protocol(args)
    if args.command == "process":
        return cmd_process(args)
    fail("Unknown command")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
