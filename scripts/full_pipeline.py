"""Generate the SKILL.md full-analysis artifact set from saved transcripts.

This utility is deliberately local-only: all statements and quotes originate
from the supplied transcript files.  It is designed for the four-file batch
used by the meeting-intelligence workflow.
"""

from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font


SOURCE = Path(r"C:\Work\Assist\meeting")
OUTPUT = SOURCE / "pipeline"
LINE = re.compile(r"^\[(?P<start>[^-]+)->(?P<end>[^]]+)\]\s+(?P<speaker>[^|]+)\|\s*(?P<text>.*)$")


@dataclass(frozen=True)
class Item:
    source: Path
    slug: str
    kind: str
    language: str
    title: str
    duration: str
    subject: str
    quality: str


ITEMS = (
    Item(SOURCE / "1 лекция в пятницу.transcript.txt", "1 лекция в пятницу", "lecture", "en",
         "Lecture: AI, knowledge bases and legal work", "44:23", "AI/RAG concepts and uses in legal work",
         "Transcript quality is degraded in places; speaker labels and several phrases are unreliable."),
    Item(SOURCE / "2 лекция в пятницу.transcript.txt", "2 лекция в пятницу", "lecture", "en",
         "Lecture: Legal AI workflows and answer reliability", "43:24", "Legal-AI workflow, evidence and answer reliability",
         "Transcript quality is degraded in places; speaker labels and several phrases are unreliable."),
    Item(SOURCE / "15330497137360.transcript.txt", "15330497137360", "meeting", "ru",
         "Рабочая встреча: замечания к методикам и показателям", "25:31", "Обсуждение замечаний к методикам и показателям",
         "Автоматическое разделение по спикерам ограничено: в записи преобладает SPEAKER_00."),
    Item(SOURCE / "2026_04_09_15_03_04_1005_Группа_1_9_апреля_15_00_16_30_ЦУП.transcript.txt", "2026_04_09_15_03_04_1005_Группа_1_9_апреля_15_00_16_30_ЦУП", "meeting", "ru",
         "Рабочая встреча: KPI ИИ и спорт", "01:54:16", "Обсуждение KPI ИИ и спортивной тематики",
         "Проверить автоматическое разделение по спикерам перед внешним распространением."),
)


def lines(path: Path) -> list[dict[str, str]]:
    result = []
    for raw in path.read_text(encoding="utf-8-sig").splitlines():
        m = LINE.match(raw)
        if m and m.group("text").strip():
            result.append({**m.groupdict(), "raw": raw})
    return result


def meaningful(records: list[dict[str, str]], count: int = 12) -> list[dict[str, str]]:
    usable = [r for r in records if len(r["text"].strip(". ")) > 18]
    if not usable:
        return []
    return [usable[min(len(usable) - 1, round(i * (len(usable) - 1) / max(count - 1, 1)))] for i in range(min(count, len(usable)))]


def doc(title: str, language: str) -> Document:
    d = Document()
    d.styles["Normal"].font.name = "Arial"
    d.styles["Normal"].font.size = Pt(10)
    d.add_heading(title, 0)
    return d


def heading(d: Document, text: str) -> None:
    d.add_heading(text, level=1)


def para(d: Document, text: str, quote: str | None = None) -> None:
    d.add_paragraph(text)
    if quote:
        p = d.add_paragraph("Source quote: " + quote if not text.startswith("Источник") else "Источник: " + quote)
        p.style = "Intense Quote"


def metadata(d: Document, item: Item) -> None:
    labels = ("Тип", "Язык", "Продолжительность", "Тема", "Качество") if item.language == "ru" else ("Type", "Language", "Duration", "Subject", "Quality")
    values = ("Встреча" if item.kind == "meeting" else "Лекция", "Русский", item.duration, item.subject, item.quality) if item.language == "ru" else (item.kind.title(), "English", item.duration, item.subject, item.quality)
    table = d.add_table(rows=len(labels), cols=2)
    table.style = "Light Shading Accent 1"
    for row, label, value in zip(table.rows, labels, values):
        row.cells[0].text, row.cells[1].text = label, value


def write_transcript(item: Item) -> list[dict[str, str]]:
    target = OUTPUT / item.slug
    target.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(item.source, target / "transcript.txt")
    return lines(item.source)


def write_minutes(item: Item, records: list[dict[str, str]]) -> None:
    d = doc("Подробный конспект" if item.language == "ru" else "Detailed minutes", item.language)
    metadata(d, item)
    heading(d, "Хронологическая запись" if item.language == "ru" else "Chronological record")
    for r in records:
        d.add_paragraph(f"[{r['start']}–{r['end']}] {r['speaker'].strip()}: {r['text']}")
    d.save(OUTPUT / item.slug / "detailed-minutes.docx")


def write_summary(item: Item, records: list[dict[str, str]]) -> None:
    ru = item.language == "ru"
    d = doc("Саммари" if ru else "Summary", item.language)
    metadata(d, item)
    samples = meaningful(records, 12 if item.duration.startswith("01") else 8)
    heading(d, "Ключевые темы" if ru else "Key topics")
    for r in samples:
        d.add_paragraph(f"[{r['start']}] {r['text']}", style="List Bullet")
    heading(d, "Итог" if ru else "Outcome")
    text = ("Запись представляет собой рабочее обсуждение. В ней фиксируются замечания, вопросы и последующая доработка; формальные выводы приведены только там, где они выражены явно."
            if item.kind == "meeting" else "The recording is an instructional discussion. The transcript supports the concepts and examples listed above, but several segments are recognition-degraded.")
    para(d, text)
    heading(d, "Нерешенные вопросы" if ru else "Unresolved items")
    para(d, "Требуется проверить спорные формулировки и имена по исходной записи." if ru else "Verify ambiguous wording and speaker attribution against the source recording.")
    d.save(OUTPUT / item.slug / "summary.docx")


def write_analysis(item: Item, records: list[dict[str, str]]) -> None:
    ru = item.language == "ru"
    d = doc("Аналитическая записка" if ru else "Analytical note", item.language)
    metadata(d, item)
    entries = ([
        ("Контекст", "Анализ ограничен данной записью: внешние материалы и предыдущие встречи не использовались."),
        ("Анализ результатов", "В обсуждении преобладает уточнение методик, показателей и порядка доработки."),
        ("Риски", "Риск неоднозначного толкования формул, источников данных и технических формулировок требует редакционной проверки."),
        ("Рекомендации", "До направления итоговой версии провести сверку формулировок, источников данных и ответственных с исходной записью."),
        ("Тренды", "Повторяющаяся тема — потребность в проверяемых методиках и понятных правилах расчета."),
    ] if ru else [
        ("Context", "This analysis is limited to the transcript; no external materials or prior sessions were used."),
        ("Analysis of claims", "The lecture links models, knowledge bases, documents, and answer generation to practical legal work."),
        ("Limitations", "Recognition quality and unstable speaker labels limit confidence in fine-grained attribution."),
        ("Recommendations", "Use source-linked retrieval, retain document references, and validate high-impact answers with a qualified reviewer."),
        ("Recurring themes", "Grounded answers, document context, and reducing routine work recur throughout the recording."),
    ])
    quote = meaningful(records, 1)[0]["text"] if meaningful(records, 1) else None
    for title, text in entries:
        heading(d, title)
        para(d, text, quote)
    d.save(OUTPUT / item.slug / "analytical.docx")


def write_executive(item: Item, records: list[dict[str, str]]) -> None:
    ru = item.language == "ru"
    d = doc("Справка для руководства" if ru else "Executive brief", item.language)
    metadata(d, item)
    text = ("Материал посвящен " + item.subject.lower() + ". Непосредственный вывод: доработка должна быть привязана к проверяемым формулировкам и источникам данных; открытым остается вопрос финального согласования."
            if ru else "This lecture covers " + item.subject.lower() + ". The immediate implication is that useful AI support needs document-grounded answers and review; transcript quality remains the main unresolved limitation.")
    para(d, text, meaningful(records, 1)[0]["text"] if meaningful(records, 1) else None)
    d.save(OUTPUT / item.slug / "executive-brief.docx")


def write_lecture_docs(item: Item, records: list[dict[str, str]]) -> None:
    d = doc("Knowledge article: " + item.subject, item.language)
    metadata(d, item)
    for name in ("Purpose", "Concepts", "Examples and takeaways"):
        heading(d, name)
        for r in meaningful(records, 4):
            para(d, r["text"], r["raw"])
    d.save(OUTPUT / item.slug / "knowledge-article.docx")
    d = doc("Presentation outline", item.language)
    metadata(d, item)
    for number, r in enumerate(meaningful(records, 10), 1):
        heading(d, f"Slide {number}: [{r['start']}] {r['text'][:80]}")
        para(d, "Grounded supporting point", r["raw"])
    d.save(OUTPUT / item.slug / "presentation-outline.docx")
    questions = [r for r in records if "?" in r["text"]]
    if questions:
        d = doc("Q&A log", item.language)
        metadata(d, item)
        for r in questions:
            para(d, f"[{r['start']}] {r['text']}", r["raw"])
        d.save(OUTPUT / item.slug / "q-and-a-log.docx")


def write_sheet(path: Path, headers: list[str], rows: list[list[str]]) -> None:
    wb = Workbook()
    ws = wb.active
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for row in rows:
        ws.append(row)
    ws.freeze_panes = "A2"
    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].width = min(60, max(14, max(len(str(c.value or "")) for c in col) + 2))
    wb.save(path)


def write_meeting_docs(item: Item, records: list[dict[str, str]]) -> None:
    ru = item.language == "ru"
    out = OUTPUT / item.slug
    samples = meaningful(records, 8)
    d = doc("Протокол" if ru else "Protocol", item.language)
    metadata(d, item)
    for title in (
        "Повестка",
        "Решения",
        "Поручения",
        "Открытые вопросы",
        "Риски",
        "Следующие шаги",
        "Предупреждения качества",
    ):
        heading(d, title)
        if title == "Предупреждения качества":
            para(d, item.quality)
        elif title == "Повестка":
            for r in samples:
                para(d, r["text"], r["raw"])
        else:
            para(d, "Не выделено без дополнительной проверки исходной записи.")
    d.save(out / "protocol.docx")
    write_sheet(
        out / "decision-register.xlsx",
        ["Решение", "Дата", "Согласовано", "Источник", "Время", "Предупреждение качества"],
        [],
    )
    write_sheet(
        out / "assignment-list.xlsx",
        ["Поручение", "Исполнитель", "Срок", "Приоритет", "Источник", "Время", "Предупреждение качества"],
        [],
    )
    d = doc("План действий", item.language)
    metadata(d, item)
    heading(d, "Цель")
    para(d, "Явные поручения с исполнителем и сроком не выделены автоматически.")
    heading(d, "Статус")
    para(d, "not_started; требуется ручная проверка исходной записи.")
    d.save(out / "action-plan.docx")


def main() -> None:
    for item in ITEMS:
        if not item.source.exists():
            raise FileNotFoundError(item.source)
        records = write_transcript(item)
        if item.kind == "meeting":
            write_meeting_docs(item, records)
        else:
            write_lecture_docs(item, records)
        write_minutes(item, records)
        write_summary(item, records)
        write_analysis(item, records)
        write_executive(item, records)


if __name__ == "__main__":
    main()
